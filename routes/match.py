from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from db.mongo import jobs_col, cvs_col, results_col, sessions_col
from services.matcher import rank_cvs, rewrite_query, build_query
import uuid

router = APIRouter(prefix="/api/match", tags=["Matching"])

# ── Start a matching session ──────────────────────────────────────────────────

class MatchRequest(BaseModel):
    job_id: str

@router.post("/")
async def start_match(req: MatchRequest):
    """
    Run the two-stage ranking pipeline for a job.
    Returns a session_id to use for feedback rounds.
    """
    # Load job
    job = await jobs_col.find_one({"job_id": req.job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Load all CVs
    cv_records = await cvs_col.find({}, {"_id": 0}).to_list(length=1000)
    if not cv_records:
        raise HTTPException(status_code=400, detail="No CVs uploaded yet")

    # Run ranking
    results = await rank_cvs(job["description"], cv_records)

    # Delete any existing sessions for this job to clear feedback history
    await sessions_col.delete_many({"job_id": req.job_id})

    # Create session for human-in-the-loop
    session_id = f"session_{uuid.uuid4().hex[:8]}"
    session    = {
        "session_id":      session_id,
        "job_id":          req.job_id,
        "jd_text":         job["description"],
        "cv_records":      cv_records,
        "round":           1,
        "history":         [],
        "current_results": results,
        "approved":        False,
    }
    await sessions_col.insert_one(session)

    # Save round 1 results to results collection
    await results_col.update_one(
        {"job_id": req.job_id},
        {"$set": {
            "job_id":   req.job_id,
            "job_title": job["title"],
            "rounds":   [{"round": 1, "results": results}],
            "final":    results,
            "approved": False,
        }},
        upsert=True,
    )

    return {
        "session_id":    session_id,
        "job_id":        req.job_id,
        "round":         1,
        "top_candidates": results[:5],
        "all_results":   results,
    }


# ── Submit feedback and re-rank ───────────────────────────────────────────────

class FeedbackRequest(BaseModel):
    feedback: str   # "yes" to approve, or text feedback to re-rank

@router.post("/{session_id}/feedback")
async def submit_feedback(session_id: str, req: FeedbackRequest):
    """
    Human-in-the-loop: submit recruiter feedback.
    Type 'yes' to approve, or give text feedback to trigger a re-rank.
    """
    session = await sessions_col.find_one({"session_id": session_id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session["approved"]:
        raise HTTPException(status_code=400, detail="Session already approved")

    current_results = session["current_results"]
    history         = session["history"]
    jd_text         = session["jd_text"]
    cv_records      = session["cv_records"]
    round_num       = session["round"]

    # Record this round in history
    top = current_results[0] if current_results else {}
    history.append({
        "round":     round_num,
        "top_cv":    top.get("cv_id", ""),
        "top_score": top.get("final_score", 0),
        "feedback":  req.feedback if req.feedback.lower() not in ("yes", "y", "approve") else None,
        "approved":  req.feedback.lower() in ("yes", "y", "approve"),
    })

    # Approved?
    if req.feedback.lower() in ("yes", "y", "approve"):
        await sessions_col.update_one(
            {"session_id": session_id},
            {"$set": {"approved": True, "history": history}},
        )
        await results_col.update_one(
            {"job_id": session["job_id"]},
            {"$set": {"final": current_results, "approved": True}},
        )
        return {
            "status":    "approved",
            "round":     round_num,
            "message":   "Ranking approved ✅",
            "final":     current_results,
        }

    # Max rounds check (3 rounds)
    if round_num >= 3:
        await sessions_col.update_one(
            {"session_id": session_id},
            {"$set": {"approved": True, "history": history}},
        )
        await results_col.update_one(
            {"job_id": session["job_id"]},
            {"$set": {"final": current_results, "approved": True}},
        )
        return {
            "status":  "max_rounds",
            "message": "Max rounds reached — using last ranking",
            "final":   current_results,
        }

    # Re-rank with rewritten query
    base_q   = build_query("ideal", "candidate", jd_text)
    new_q    = await rewrite_query(jd_text, base_q, req.feedback, history)
    query_override = {cv["cv_id"]: new_q for cv in cv_records}

    new_results = await rank_cvs(jd_text, cv_records, query_override=query_override)
    new_round   = round_num + 1

    # Save new round
    await sessions_col.update_one(
        {"session_id": session_id},
        {"$set": {
            "round":           new_round,
            "current_results": new_results,
            "history":         history,
        }},
    )
    await results_col.update_one(
        {"job_id": session["job_id"]},
        {"$push": {"rounds": {"round": new_round, "results": new_results}},
         "$set":  {"final": new_results}},
    )

    return {
        "status":          "re_ranked",
        "round":           new_round,
        "rewritten_query": new_q,
        "top_candidates":  new_results[:5],
        "all_results":     new_results,
    }


# ── Get results for a job ─────────────────────────────────────────────────────

@router.get("/results/{job_id}")
async def get_results(job_id: str):
    """Get the final ranking results for a job."""
    result = await results_col.find_one({"job_id": job_id}, {"_id": 0})
    if not result:
        raise HTTPException(status_code=404, detail="No results found for this job")
    return result


# ── AI CV Processing & Agentic Ranking Endpoints ──────────────────────────────

import os
import fitz # PyMuPDF
import docx as docx_lib
from urllib.parse import urlparse
from config import UPLOAD_DIR
from db.mongo import db

class ProcessCVRequest(BaseModel):
    cv_url: str
    application_id: str
    job_id: str
    job_title: str
    job_description: str
    job_skills: list

class AgenticRankRequest(BaseModel):
    job_id: str
    job_title: str
    job_description: str
    job_skills: list
    recruiter_query: str
    round: int

def extract_text_from_path(path: str) -> str:
    if path.endswith(".docx"):
        try:
            doc  = docx_lib.Document(path)
            text = " ".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += " " + cell.text
            return text
        except Exception as e:
            print("Docx extract error:", e)
            return ""
    elif path.endswith(".pdf"):
        try:
            doc  = fitz.open(path)
            text = " ".join(page.get_text() for page in doc)
            doc.close()
            return text
        except Exception as e:
            print("PDF extract error:", e)
            return ""
    return ""

@router.post("/process")
async def process_cv(req: ProcessCVRequest):
    parsed_url = urlparse(req.cv_url)
    filename = os.path.basename(parsed_url.path)
    local_path = os.path.join(UPLOAD_DIR, filename)
    
    cv_text = ""
    if os.path.exists(local_path):
        cv_text = extract_text_from_path(local_path)
    
    if not cv_text:
        cv_text = f"Resume of applicant. Skills: {', '.join(req.job_skills) if req.job_skills else 'Software Engineering'}."
        
    from services.matcher import llm_score, hybrid_score, generate_candidate_feedback
    from services.embeddings import hf_embedding_func
    from sklearn.metrics.pairwise import cosine_similarity
    
    try:
        jd_emb = await hf_embedding_func([req.job_description])
        cv_emb = await hf_embedding_func([cv_text])
        sim = cosine_similarity(jd_emb, cv_emb)[0][0]
        sim_score = round(max(0.0, float(sim)) * 100, 1)
    except Exception as e:
        print("Embedding similarity error:", e)
        sim_score = 50.0
        
    llm_total, verdict, scores, reasons = await llm_score(req.job_description, cv_text)
    final_score = hybrid_score(sim_score, llm_total)
    
    feedback_text = await generate_candidate_feedback(
        match_score=final_score,
        job_title=req.job_title,
        job_description=req.job_description,
        cv_summary=cv_text
    )
    
    application_col = db["applications"]
    await application_col.update_one(
        {"$or": [{"id": req.application_id}, {"application_id": req.application_id}]},
        {"$set": {
            "status": "processed",
            "match_score": final_score,
            "feedback": feedback_text,
            "rag_results": {
                "feedback": verdict,
                "ranking_reason": reasons.get("explanation") or verdict,
                "scores": scores,
                "reasons": reasons
            }
        }}
    )
    
    return {
        "success": True,
        "match_score": final_score,
        "verdict": verdict,
        "scores": scores
    }

@router.post("/agentic")
async def agentic_rank(req: AgenticRankRequest):
    application_col = db["applications"]
    apps = await application_col.find({"job_id": req.job_id}).to_list(length=1000)
    if not apps:
        raise HTTPException(status_code=400, detail="No applications found for this job")
        
    cv_records = []
    for app in apps:
        cv_url = app.get("cv_url", "")
        if not cv_url:
            continue
            
        parsed_url = urlparse(cv_url)
        filename = os.path.basename(parsed_url.path)
        local_path = os.path.join(UPLOAD_DIR, filename)
        
        cv_text = ""
        if os.path.exists(local_path):
            cv_text = extract_text_from_path(local_path)
            
        if not cv_text:
            cv_text = f"Resume of {app.get('candidate_name', 'Unknown')}. Skills: {', '.join(app.get('skills', [])) or 'Development'}."
            
        cv_records.append({
            "cv_id": app.get("id"),
            "text": cv_text,
            "category": "General",
            "original_filename": filename
        })
        
    if not cv_records:
        raise HTTPException(status_code=400, detail="No applications with CV files found")
        
    query_override = None
    if req.recruiter_query:
        query_override = {cv["cv_id"]: req.recruiter_query for cv in cv_records}
        
    ranked_results = await rank_cvs(req.job_description, cv_records, query_override=query_override)
    
    from services.matcher import generate_candidate_feedback
    for row in ranked_results:
        app_id = row["cv_id"]
        final_score = row["final_score"]
        verdict = row["verdict"]
        scores = row["scores"]
        
        cv_rec = next((c for c in cv_records if c["cv_id"] == app_id), {})
        cv_text = cv_rec.get("text", "")
        
        feedback_text = await generate_candidate_feedback(
            match_score=final_score,
            job_title=req.job_title,
            job_description=req.job_description,
            cv_summary=cv_text
        )
        
        await application_col.update_one(
            {"id": app_id},
            {"$set": {
                "status": "processed",
                "match_score": final_score,
                "feedback": feedback_text,
                "rag_results": {
                    "feedback": verdict,
                    "ranking_reason": verdict,
                    "scores": scores
                }
            }}
        )
        
    return {
        "success": True,
        "ranked": len(ranked_results)
    }

@router.post("/{job_id}/rank-and-feedback")
async def rank_and_feedback(job_id: str):
    job = await db["jobs"].find_one({"$or": [{"id": job_id}, {"job_id": job_id}]})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
        
    application_col = db["applications"]
    apps = await application_col.find({"job_id": job_id}).to_list(length=1000)
    if not apps:
        raise HTTPException(status_code=400, detail="No applications found for this job")
        
    await application_col.update_many({"job_id": job_id}, {"$set": {"feedback": ""}})
    
    cv_records = []
    for app in apps:
        cv_url = app.get("cv_url", "")
        if not cv_url:
            continue
            
        parsed_url = urlparse(cv_url)
        filename = os.path.basename(parsed_url.path)
        local_path = os.path.join(UPLOAD_DIR, filename)
        
        cv_text = ""
        if os.path.exists(local_path):
            cv_text = extract_text_from_path(local_path)
            
        if not cv_text:
            cv_text = f"Resume of {app.get('candidate_name', 'Unknown')}. Skills: {', '.join(app.get('skills', [])) or 'Development'}."
            
        cv_records.append({
            "cv_id": app.get("id"),
            "text": cv_text,
            "category": "General",
            "original_filename": filename
        })
        
    ranked_results = []
    if cv_records:
        ranked_results = await rank_cvs(job["description"], cv_records)
        
    def format_skills(skills):
        if not skills:
            return "your field"
        skills = [s.strip() for s in skills if s and s.strip()]
        if not skills:
            return "your field"
        if len(skills) == 1:
            return skills[0]
        elif len(skills) == 2:
            return f"{skills[0]} and {skills[1]}"
        else:
            first_part = ", ".join(skills[:2])
            return f"{first_part}, and {skills[2]}"
            
    for app in apps:
        app_id = app.get("id")
        
        row = next((r for r in ranked_results if r["cv_id"] == app_id), None)
        final_score = row["final_score"] if row else app.get("match_score")
        verdict = row["verdict"] if row else None
        scores = row["scores"] if row else None
        
        skills_list = []
        c_email = app.get("candidate_email")
        if c_email:
            cand = await db["candidates"].find_one({"email": c_email.strip().lower()})
            if cand:
                skills_list = cand.get("skills") or cand.get("extracted_skills") or []
        if not skills_list:
            skills_list = app.get("skills") or app.get("extracted_skills") or []
            
        status = app.get("status", "pending") or "pending"
        status = status.lower().strip()
        
        if status == "shortlisted":
            skills_str = format_skills(skills_list)
            feedback = f"Congratulations! Your application for {job['title']} has been shortlisted. Your skills in {skills_str} align well with our requirements. Our team will contact you shortly with next steps."
        elif status == "rejected":
            feedback = f"Thank you for your application for {job['title']}. While your background is impressive, we have decided to move forward with other candidates whose experience more closely matches our needs. We wish you the best in your job search."
        else:
            feedback = f"Your application for {job['title']} is currently under review. We will notify you once a decision has been made."
            
        update_doc = {
            "feedback": feedback
        }
        if final_score is not None:
            update_doc["match_score"] = final_score
        if verdict is not None:
            explanation = row.get("reasons", {}).get("explanation", "") if row else ""
            update_doc["rag_results"] = {
                "feedback": verdict,
                "ranking_reason": explanation or verdict,
                "scores": scores,
                "reasons": row.get("reasons", {}) if row else {}
            }
            
        await application_col.update_one({"id": app_id}, {"$set": update_doc})
        
    return {
        "success": True,
        "message": f"Ranking completed and feedback generated for {len(apps)} candidates."
    }


