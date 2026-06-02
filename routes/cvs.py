import os
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Optional
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Header
from fastapi.responses import JSONResponse, FileResponse
from db.mongo import cvs_col, db
from config import UPLOAD_DIR
from routes.auth import get_user_by_token
import fitz
import docx as docx_lib

router = APIRouter(prefix="/api/cvs", tags=["CVs"])

def get_text(path: str) -> str:
    if path.endswith(".docx"):
        doc  = docx_lib.Document(path)
        text = " ".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += " " + cell.text
        return text
    elif path.endswith(".pdf"):
        doc  = fitz.open(path)
        text = " ".join(page.get_text() for page in doc)
        doc.close()
        return text
    return ""

@router.post("/upload")
async def upload_cvs(files: list[UploadFile] = File(...)):
    """
    Upload one or more CV files (.pdf, .docx) or a single .zip containing them.
    Extracted text is saved to MongoDB.
    """
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    saved = []

    for upload in files:
        content  = await upload.read()
        filename = upload.filename.lower()

        if filename.endswith(".zip"):
            # Extract zip and process each CV inside
            with tempfile.TemporaryDirectory() as tmp:
                zip_path = os.path.join(tmp, "upload.zip")
                with open(zip_path, "wb") as f:
                    f.write(content)
                with zipfile.ZipFile(zip_path, "r") as z:
                    z.extractall(tmp)
                for fp in Path(tmp).rglob("*"):
                    if fp.is_file() and fp.suffix.lower() in {".docx", ".pdf"}:
                        dest = os.path.join(UPLOAD_DIR, fp.name)
                        shutil.copy2(fp, dest)
                        text    = get_text(dest)
                        cv_id   = fp.stem
                        cv_data = {
                            "cv_id":             cv_id,
                            "original_filename": fp.name,
                            "category":          "General",
                            "text":              text,
                            "path":              dest,
                        }
                        await cvs_col.update_one(
                            {"cv_id": cv_id},
                            {"$set": cv_data},
                            upsert=True,
                        )
                        saved.append(cv_id)
        elif filename.endswith((".pdf", ".docx")):
            dest  = os.path.join(UPLOAD_DIR, upload.filename)
            with open(dest, "wb") as f:
                f.write(content)
            text    = get_text(dest)
            cv_id   = Path(upload.filename).stem
            cv_data = {
                "cv_id":             cv_id,
                "original_filename": upload.filename,
                "category":          "General",
                "text":              text,
                "path":              dest,
            }
            await cvs_col.update_one(
                {"cv_id": cv_id},
                {"$set": cv_data},
                upsert=True,
            )
            saved.append(cv_id)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {upload.filename}")

    return {"message": f"Uploaded {len(saved)} CV(s)", "cv_ids": saved}


@router.get("/")
async def list_cvs():
    """Return all CVs stored in MongoDB."""
    cvs = await cvs_col.find({}, {"_id": 0, "text": 0}).to_list(length=1000)
    return {"count": len(cvs), "cvs": cvs}


@router.delete("/{cv_id}")
async def delete_cv(cv_id: str):
    result = await cvs_col.delete_one({"cv_id": cv_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="CV not found")
    return {"message": f"Deleted CV: {cv_id}"}


candidate_cv_router = APIRouter(prefix="/api/candidate", tags=["Candidate CV"])

@candidate_cv_router.delete("/cv")
async def delete_candidate_cv(user: dict = Depends(get_user_by_token)):
    if user.get("role") != "candidate":
        raise HTTPException(status_code=403, detail="Only candidates can delete their CV")
        
    user_id = user.get("id")
    
    # 1. Lookup the candidate's CV document
    cv_record = await cvs_col.find_one({"user_id": user_id})
    if cv_record:
        file_path = cv_record.get("path")
        # Delete local file from disk
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                print(f"Failed to remove file from disk: {e}")
        
        # Delete the CV record from cvs collection
        await cvs_col.delete_one({"user_id": user_id})
        
    # 2. Update candidates collection (remove fields)
    candidate_update_result = await db["candidates"].update_one(
        {"user_id": user_id},
        {"$unset": {"cv_url": "", "cv_filename": "", "cv_uploaded_at": "", "cv_upload_date": ""}}
    )
    
    # If not matched by user_id, fallback to search by user email
    if candidate_update_result.matched_count == 0:
        email = user.get("email")
        if email:
            await db["candidates"].update_one(
                {"email": email.strip().lower()},
                {"$unset": {"cv_url": "", "cv_filename": "", "cv_uploaded_at": "", "cv_upload_date": ""}}
            )
            
    return {"message": "CV deleted successfully"}


@router.get("/{cv_id}/download")
async def download_cv(
    cv_id: str, 
    authorization: Optional[str] = Header(None),
    token: Optional[str] = None
):
    """
    Download a CV by its cv_id, adding debug logging about file existence,
    paths, database registration, and requester role.
    """
    # 1. Retrieve user details if authenticated
    user = None
    auth_token = None
    if authorization and authorization.startswith("Bearer "):
        auth_token = authorization.split(" ")[1]
    elif token:
        auth_token = token

    if auth_token:
        user = await db["users"].find_one({"id": auth_token}, {"_id": 0})

    user_permission = user.get("role") if user else "Anonymous (No Auth Token)"
    user_email = user.get("email") if user else "Anonymous"

    # 2. Database existence check (in cvs & applications collections)
    cv_record = await cvs_col.find_one({"cv_id": cv_id})
    db_exists = cv_record is not None
    filename = None
    
    if db_exists:
        filename = cv_record.get("filename") or cv_record.get("original_filename")
    else:
        # Fallback: check if cv_id is a user_id or username
        cv_record = await cvs_col.find_one({"user_id": cv_id})
        if cv_record:
            db_exists = True
            filename = cv_record.get("filename") or cv_record.get("original_filename")
        else:
            # Check by application_id in applications collection
            app_record = await db["applications"].find_one({"id": cv_id})
            if app_record:
                cv_url = app_record.get("cv_url", "")
                if cv_url:
                    filename = cv_url.split("/")[-1]
                    db_exists = True

    # 3. Storage existence check
    target_path = ""
    file_exists_on_disk = False
    if db_exists and filename:
        target_path = os.path.join(UPLOAD_DIR, filename)
        file_exists_on_disk = os.path.exists(target_path)

    # 4. Detailed Logging
    print("\n" + "="*60)
    print("DEBUG LOG: CV DOWNLOAD REQUEST (by cv_id)")
    print(f"Requested CV ID / Reference: {cv_id}")
    print(f"Target Filename: {filename or 'N/A'}")
    print(f"Target Path: {target_path or 'N/A'}")
    print(f"Exists in Database: {db_exists}")
    print(f"Exists on Disk: {file_exists_on_disk}")
    print(f"Requester Email: {user_email}")
    print(f"Requester Role/Permission Level: {user_permission}")
    print("="*60 + "\n")

    if not db_exists or not file_exists_on_disk:
        raise HTTPException(status_code=404, detail="CV not found")

    return FileResponse(target_path, filename=filename)


