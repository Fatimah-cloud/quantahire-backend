# routes/upload.py

"""
routes/upload.py

Key addition: after saving a CV, if job_id is provided,
we trigger async RAG indexing into rag_storage/{job_id}/.
"""

import os
import uuid
import asyncio
from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Header
from typing import Optional
from config import UPLOAD_DIR
from db.mongo import db, cvs_col
from datetime import datetime
import fitz
import docx as docx_lib

router = APIRouter(prefix="/api/upload", tags=["Uploads"])


def extract_text(path: str) -> str:
    if path.endswith(".docx"):
        try:
            doc  = docx_lib.Document(path)
            text = " ".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += " " + cell.text
            return text
        except Exception as e:
            print(f"[upload] DOCX extract error: {e}")
            return ""
    elif path.endswith(".pdf"):
        try:
            doc  = fitz.open(path)
            text = " ".join(page.get_text() for page in doc)
            doc.close()
            return text
        except Exception as e:
            print(f"[upload] PDF extract error: {e}")
            return ""
    return ""


async def _index_in_background(cv_path: str, job_id: str):
    """Fire-and-forget RAG indexing so the upload response is not delayed."""
    try:
        from services.matcher import index_cv_into_rag
        await index_cv_into_rag(cv_path, job_id)
    except Exception as e:
        print(f"[upload] Background RAG indexing failed for {cv_path}: {e}")


@router.post("/")
async def upload_file(
    file:          UploadFile    = File(...),
    user_id:       Optional[str] = Form(None),
    job_id:        Optional[str] = Form(None),
    authorization: Optional[str] = Header(None),
):
    os.makedirs(UPLOAD_DIR, exist_ok=True)

    resolved_user_id = user_id
    if not resolved_user_id and authorization and authorization.startswith("Bearer "):
        resolved_user_id = authorization.split(" ")[1]

    unique_prefix  = uuid.uuid4().hex[:8]
    sanitized_name = f"{unique_prefix}_{file.filename.replace(' ', '_')}"
    dest_path      = os.path.join(UPLOAD_DIR, sanitized_name)

    try:
        content = await file.read()
        with open(dest_path, "wb") as f:
            f.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

    file_url = f"http://localhost:8000/uploads/{sanitized_name}"

    if resolved_user_id:
        extracted_text = ""
        if file.filename.lower().endswith((".pdf", ".docx")):
            extracted_text = extract_text(dest_path)

        cv_id   = f"cv_{uuid.uuid4().hex[:8]}"
        cv_data = {
            "cv_id":             cv_id,
            "user_id":           resolved_user_id,
            "job_id":            job_id or "",
            "original_filename": file.filename,
            "filename":          sanitized_name,
            "file_url":          file_url,
            "path":              dest_path,
            "text":              extracted_text,
            "category":          "General",
            "created_at":        datetime.utcnow().isoformat(),
        }

        await cvs_col.update_one(
            {"user_id": resolved_user_id},
            {"$set": cv_data},
            upsert=True,
        )

        candidate = await db["candidates"].find_one({"user_id": resolved_user_id})
        if candidate:
            await db["candidates"].update_one(
                {"user_id": resolved_user_id},
                {"$set": {
                    "cv_url":         file_url,
                    "cv_filename":    file.filename,
                    "cv_uploaded_at": datetime.utcnow().isoformat(),
                }},
            )
        else:
            user = await db["users"].find_one({"id": resolved_user_id})
            if user:
                email_lower = user.get("email", "").strip().lower()
                if email_lower:
                    await db["candidates"].update_one(
                        {"email": email_lower},
                        {"$set": {
                            "cv_url":         file_url,
                            "cv_filename":    file.filename,
                            "cv_uploaded_at": datetime.utcnow().isoformat(),
                        }},
                    )

        if job_id:
            user_doc   = await db["users"].find_one({"id": resolved_user_id})
            cand_doc   = candidate or await db["candidates"].find_one({"user_id": resolved_user_id})
            cand_email = (user_doc or {}).get("email") or (cand_doc or {}).get("email") or ""
            cand_name  = (user_doc or {}).get("full_name") or (cand_doc or {}).get("full_name") or "Unknown"

            existing_app = await db["applications"].find_one({
                "job_id":          job_id,
                "candidate_email": cand_email,
            })

            if existing_app:
                await db["applications"].update_one(
                    {"_id": existing_app["_id"]},
                    {"$set": {
                        "cv_url":     file_url,
                        "status":     "pending",
                        "updated_at": datetime.utcnow().isoformat(),
                    }},
                )
            else:
                app_id  = f"app_{uuid.uuid4().hex[:8]}"
                app_doc = {
                    "id":              app_id,
                    "application_id":  app_id,
                    "job_id":          job_id,
                    "candidate_id":    resolved_user_id,
                    "candidate_email": cand_email,
                    "candidate_name":  cand_name,
                    "cv_url":          file_url,
                    "status":          "pending",
                    "match_score":     None,
                    "feedback":        "",
                    "rag_results":     {},
                    "created_date":    datetime.utcnow().isoformat(),
                    "created_at":      datetime.utcnow().isoformat(),
                }
                await db["applications"].insert_one(app_doc)

            # Trigger RAG indexing as a background task — does not block the response
            asyncio.create_task(_index_in_background(dest_path, job_id))
            print(f"[upload] RAG indexing scheduled for {file.filename} in job graph '{job_id}'")

    return {
        "message":           "File uploaded successfully",
        "filename":          sanitized_name,
        "original_filename": file.filename,
        "file_url":          file_url,
        "rag_indexing":      "scheduled" if job_id else "skipped (no job_id)",
    }
